# import os
# from pdf2image import convert_from_path
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from tenacity import retry, stop_after_attempt, wait_exponential
# import logging
# import json
# import pytesseract
# from pdf2image.exceptions import PDFPageCountError
# from reportlab.lib import colors

# # Import shared objects from config.py
# from config import gemini_model, supabase, SUPABASE_STORAGE_BUCKET, SUPABASE_REPORT_PATH_PREFIX, redis_conn

# # Set up logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # Function to log progress and store in Redis
# def log_progress(job_id: str, step: str, message: str, data: dict = None):
#     """Log a progress step and store it in Redis."""
#     log_entry = {
#         "step": step,
#         "message": message,
#         "data": data or {},
#         "timestamp": str(os.times()[4])  # System time
#     }
#     logger.info(f"Job {job_id} - {step}: {message}")
#     redis_conn.rpush(f"job_logs:{job_id}", json.dumps(log_entry))

# # Helper function to download resume from Supabase Storage
# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def download_resume(resume_path: str) -> str:
#     try:
#         response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).download(resume_path)
#         local_path = f"/tmp/{os.path.basename(resume_path)}"
#         with open(local_path, "wb") as f:
#             f.write(response)
#         return local_path
#     except Exception as e:
#         raise Exception(f"Failed to download resume from Supabase: {str(e)}")

# # Helper function to get the total number of pages in the PDF
# def get_pdf_page_count(pdf_path: str) -> int:
#     try:
#         from pdf2image import pdfinfo_from_path
#         pdf_info = pdfinfo_from_path(pdf_path)
#         return int(pdf_info["Pages"])
#     except PDFPageCountError as e:
#         raise Exception(f"Failed to determine PDF page count: {str(e)}")

# # Helper function to extract text from PDF using OCR, processing one page at a time
# def extract_text_from_pdf(pdf_path: str, job_id: str) -> str:
#     try:
#         total_pages = get_pdf_page_count(pdf_path)
#         log_progress(job_id, "extract_text", f"PDF has {total_pages} pages")

#         text = ""
#         custom_config = r'--oem 3 --psm 6'  # OEM 3 (default), PSM 6 (assume a single uniform block of text)

#         for page_num in range(1, total_pages + 1):
#             log_progress(job_id, "extract_text", f"Converting page {page_num}/{total_pages} to image at 200 DPI")
#             images = convert_from_path(pdf_path, dpi=200, first_page=page_num, last_page=page_num)
#             if not images:
#                 log_progress(job_id, "extract_text", f"No image generated for page {page_num}")
#                 continue

#             page_text = pytesseract.image_to_string(images[0], config=custom_config)
#             text += f"\n--- Page {page_num} ---\n{page_text}"
#             log_progress(job_id, "extract_text", f"Extracted text from page {page_num}/{total_pages}", {
#                 "page_text_length": len(page_text),
#                 "page_text": page_text[:500]  # Log first 500 chars for debugging (avoid Redis size limits)
#             })

#             del images

#         log_progress(job_id, "extract_text", "Text extracted successfully", {
#             "text_length": len(text),
#             "full_text": text[:1000]  # Log first 1000 chars of full text for debugging
#         })
#         return text.strip() or "No text extracted"
#     except Exception as e:
#         raise Exception(f"Failed to extract text from PDF: {str(e)}")

# # Helper function to clean Gemini AI output by removing Markdown headings
# def clean_gemini_output(text: str) -> str:
#     lines = text.split("\n")
#     cleaned_lines = [line for line in lines if not line.strip().startswith("##")]
#     return "\n".join(cleaned_lines).strip()

# # ... [previous imports and unchanged functions remain the same] ...

# # Helper function to generate the overall report using a single Gemini API call
# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def generate_report(resume_text: str, job_description: str) -> dict:
#     try:
#         # Log the input to Gemini for debugging
#         log_progress("debug", "generate_report_input", "Sending to Gemini", {
#             "resume_text": resume_text[:1000],
#             "job_description": job_description[:1000]
#         })

#         prompt = f"""
#         Analyze the following resume against the job description and provide a detailed suitability report for HR staffing purposes. The report should be structured, clear, and concise, with no asterisks (**) around headings or content. Focus on validating the candidate’s skills, experience, and education with high specificity—e.g., identify exact skills used in named projects (like 'Robotic Mechanic') and contributions (like 'Implemented servo control with Python'). Explain the scoring rationale for each section in detail, tying it to specific evidence from the resume and job description. Ensure the output fits within an A4 page context and suggest precise, actionable ways to enhance the candidate’s score. Follow these instructions precisely:

#         1. **Candidate Details**:
#            - Extract: Name, Phone Number, Email, LinkedIn URL, GitHub URL from the resume. If not found, write "Not provided".

#         2. **Sectional Analysis (Skills, Work Experience, Relevant Projects, Education)**:
#            - Matches: List specific items from the job description that match the resume, with detailed validation (e.g., "Java - Listed under skills and used in 'Robotic Mechanic' project for backend logic"). Include examples where possible.
#            - Non-Matches: List key items from the job description missing in the resume, with specific explanation (e.g., "AWS - Required for cloud deployment, but no mention of cloud experience in resume or projects").
#            - Summary: Provide a detailed, specific summary of alignment, including strengths (e.g., "Vinay’s Java expertise shines in Robotic Mechanic") and gaps (e.g., "Lacks cloud skills critical for this role"). Avoid vague phrases like "good fit" or "some experience"—use precise evidence.
#            - Score: Assign a score out of 100 based on relevance and completeness:
#              - Skills: 100 if all key skills match with evidence of use (e.g., in projects), 0 if none match, proportional otherwise (e.g., 75 if 3/4 skills match with examples).
#              - Work Experience: 100 if experience fully aligns in duration and tech (e.g., "6 years in Java aligns with 6-9 years required"), 0 if irrelevant (e.g., "Only sales experience").
#              - Relevant Projects: 100 if projects directly match job needs with named examples (e.g., "Robotic Mechanic used Java and SQL"), 0 if none are relevant.
#              - Education: 100 if fully meets requirements (e.g., "B.Tech in Mechanical Engineering matches robotics focus"), 0 if irrelevant.
#            - Enhancement Tips: Suggest specific, actionable improvements (e.g., "Gain AWS certification and add a cloud-based project to GitHub").
#            - Example: For a Java Developer JD requiring Java, Spring, and 3+ years:
#              - Matches: "Java - Used in 'Robotic Mechanic' for control logic; Spring - Implemented in 'E-commerce App' backend."
#              - Non-Matches: "AWS - No cloud experience mentioned."
#              - Summary: "Vinay’s 4 years of Java and Spring experience in named projects like Robotic Mechanic align well, but AWS is a gap."
#              - Score: "75/100 - Strong Java and Spring, missing AWS."
#              - Tips: "Complete an AWS project and list it."

#         3. **Overall Summary and Score**:
#            - Summary: Summarize the candidate’s suitability with specific highlights (e.g., "Vinay’s 4-year Java experience in Robotic Mechanic is a strength, but missing AWS limits cloud suitability") and gaps from an HR perspective. Keep it 2-3 sentences, detailed, and evidence-based.
#            - Overall Score: Calculate as a weighted average:
#              - Skills (30%) + Work Experience (20%) + Relevant Projects (30%) + Education (20%).
#              - Formula: (Skills_Score * 0.3) + (Work_Exp_Score * 0.2) + (Projects_Score * 0.3) + (Education_Score * 0.2).
#              - Return the score as a float with one decimal place (e.g., 63.5).
#            - Example: "Vinay’s Java and Spring skills in Robotic Mechanic are strong, but lack of AWS and limited project diversity lower his fit. Overall Score: 72.5/100."

#         Resume: {resume_text}
#         Job Description: {job_description}

#         Return the report in this exact format:
#         Candidate Details:
#         - Name: [Extracted Name]
#         - Phone Number: [Extracted Phone]
#         - Email: [Extracted Email]
#         - LinkedIn: [Extracted LinkedIn or Not provided]
#         - GitHub: [Extracted GitHub or Not provided]

#         Skill Match:
#         - Matches: [List with detailed validation, e.g., "Python - Used in 'Robotic Mechanic' for servo control"]
#         - Non-Matches: [List with specific explanation, e.g., "Kubernetes - No containerization experience noted"]
#         - Summary: [Detailed, evidence-based summary, e.g., "Vinay’s Python and Java skills from Robotic Mechanic align with automation needs, but Kubernetes is missing"]
#         - Score: [Score]/100 - [Detailed rationale, e.g., "80/100 - Matches 4/5 skills with project evidence, missing Kubernetes"]
#         - Enhancement Tips: [Specific suggestions, e.g., "Learn Kubernetes via a Udemy course and deploy a sample project"]

#         Work Experience:
#         - Matches: [List with details, e.g., "4 years at XYZ Corp - Worked on 'Robotic Mechanic' using Java, matching 3+ year requirement"]
#         - Non-Matches: [List with explanation, e.g., "No team lead experience - JD requires leadership, not mentioned"]
#         - Summary: [Detailed summary, e.g., "Vinay’s 4 years at XYZ Corp on Robotic Mechanic shows strong technical fit, but lacks leadership experience"]
#         - Score: [Score]/100 - [Rationale, e.g., "85/100 - Exceeds duration, missing leadership"]
#         - Enhancement Tips: [Suggestions, e.g., "Take on a team lead role in next project"]

#         Relevant Projects:
#         - Matches: [List with details, e.g., "Robotic Mechanic - Implemented Java and SQL for control systems, matches JD’s automation focus"]
#         - Non-Matches: [List with explanation, e.g., "No cloud projects - JD requires AWS experience"]
#         - Summary: [Detailed summary, e.g., "Robotic Mechanic demonstrates relevant Java and SQL skills, but no cloud projects limit broader fit"]
#         - Score: [Score]/100 - [Rationale, e.g., "70/100 - One strong match, missing cloud scope"]
#         - Enhancement Tips: [Suggestions, e.g., "Build an AWS-hosted automation project"]

#         Education:
#         - Matches: [List with details, e.g., "B.Tech in Mechanical Engineering - Matches robotics JD requirement"]
#         - Non-Matches: [List with explanation, e.g., "No advanced degree - JD prefers M.Tech, not required"]
#         - Summary: [Detailed summary, e.g., "Vinay’s B.Tech aligns with robotics focus, though an M.Tech could enhance credentials"]
#         - Score: [Score]/100 - [Rationale, e.g., "90/100 - Meets core requirement, lacks preferred advanced degree"]
#         - Enhancement Tips: [Suggestions, e.g., "Pursue an M.Tech or relevant certification"]

#         Overall Summary:
#         - [Concise, specific summary, e.g., "Vinay’s Java and SQL expertise in Robotic Mechanic suits the role, but gaps in AWS and leadership reduce fit"]
#         - Overall Score: [Weighted Score]/100
#         """

#         response = gemini_model.generate_content(prompt)
#         gemini_output = clean_gemini_output(response.text)

#         # Log Gemini's raw output for debugging
#         log_progress("debug", "generate_report_output", "Gemini response", {
#             "raw_output": gemini_output[:1000]
#         })

#         # Initialize the report with default values
#         report = {
#             "Candidate Details": {
#                 "Name": "Not provided",
#                 "Phone Number": "Not provided",
#                 "Email": "Not provided",
#                 "LinkedIn": "Not provided",
#                 "GitHub": "Not provided"
#             },
#             "Skill Match": {
#                 "Matches": "Not found",
#                 "Non-Matches": "Not found",
#                 "Summary": "Not found",
#                 "Score": 0,
#                 "Enhancement Tips": "Not found"
#             },
#             "Work Experience": {
#                 "Matches": "Not found",
#                 "Non-Matches": "Not found",
#                 "Summary": "Not found",
#                 "Score": 0,
#                 "Enhancement Tips": "Not found"
#             },
#             "Relevant Projects": {
#                 "Matches": "Not found",
#                 "Non-Matches": "Not found",
#                 "Summary": "Not found",
#                 "Score": 0,
#                 "Enhancement Tips": "Not found"
#             },
#             "Education": {
#                 "Matches": "Not found",
#                 "Non-Matches": "Not found",
#                 "Summary": "Not found",
#                 "Score": 0,
#                 "Enhancement Tips": "Not found"
#             },
#             "Overall Summary": "Not found",
#             "Overall Score": 0.0
#         }

#         # Parse the Gemini response
#         current_section = None
#         matches_lines = []
#         non_matches_lines = []
#         summary_lines = []
#         enhancement_tips_lines = []

#         for line in gemini_output.split("\n"):
#             line = line.strip()
#             if not line:
#                 continue

#             # Section headers
#             if "Candidate Details:" in line:
#                 current_section = "Candidate Details"
#             elif "Skill Match:" in line:
#                 current_section = "Skill Match"
#                 matches_lines, non_matches_lines, summary_lines, enhancement_tips_lines = [], [], [], []
#             elif "Work Experience:" in line:
#                 current_section = "Work Experience"
#                 matches_lines, non_matches_lines, summary_lines, enhancement_tips_lines = [], [], [], []
#             elif "Relevant Projects:" in line:
#                 current_section = "Relevant Projects"
#                 matches_lines, non_matches_lines, summary_lines, enhancement_tips_lines = [], [], [], []
#             elif "Education:" in line:
#                 current_section = "Education"
#                 matches_lines, non_matches_lines, summary_lines, enhancement_tips_lines = [], [], [], []
#             elif "Overall Summary:" in line:
#                 current_section = "Overall Summary"
#                 summary_lines = []
#             elif "Overall Score:" in line:
#                 score_str = line.split("Overall Score:", 1)[1].strip().split("/")[0].strip()
#                 report["Overall Score"] = float(score_str) if score_str.replace(".", "").isdigit() else 0.0

#             # Parse section content
#             elif current_section == "Candidate Details":
#                 for key in report["Candidate Details"].keys():
#                     field = f"- {key}:"
#                     if field in line:
#                         report["Candidate Details"][key] = line.split(field, 1)[1].strip()
#                         break
#             elif current_section in ["Skill Match", "Work Experience", "Relevant Projects", "Education"]:
#                 if "Matches:" in line:
#                     matches_lines.append(line.split("Matches:", 1)[1].strip())
#                 elif "Non-Matches:" in line:
#                     non_matches_lines.append(line.split("Non-Matches:", 1)[1].strip())
#                 elif "Summary:" in line:
#                     summary_lines.append(line.split("Summary:", 1)[1].strip())
#                 elif "Score:" in line:
#                     score_str = line.split("Score:", 1)[1].strip().split("/")[0].strip()
#                     report[current_section]["Score"] = int(score_str) if score_str.isdigit() else 0
#                 elif "Enhancement Tips:" in line:
#                     enhancement_tips_lines.append(line.split("Enhancement Tips:", 1)[1].strip())
#                 elif matches_lines and "Non-Matches:" not in line and "Summary:" not in line:
#                     matches_lines.append(line.strip())
#                 elif non_matches_lines and "Summary:" not in line and "Score:" not in line:
#                     non_matches_lines.append(line.strip())
#                 elif summary_lines and "Score:" not in line and "Enhancement Tips:" not in line:
#                     summary_lines.append(line.strip())
#                 elif enhancement_tips_lines and "Score:" not in line:
#                     enhancement_tips_lines.append(line.strip())

#                 report[current_section]["Matches"] = " ".join(matches_lines).strip()
#                 report[current_section]["Non-Matches"] = " ".join(non_matches_lines).strip()
#                 report[current_section]["Summary"] = " ".join(summary_lines).strip()
#                 report[current_section]["Enhancement Tips"] = " ".join(enhancement_tips_lines).strip()
#             elif current_section == "Overall Summary" and "Overall Score:" not in line:
#                 summary_lines.append(line.replace("Overall Summary:", "").strip())
#                 report["Overall Summary"] = " ".join(summary_lines).strip()

#         # Log the parsed report for debugging
#         log_progress("debug", "generate_report_parsed", "Parsed report", {
#             "report": report
#         })

#         return report

#     except Exception as e:
#         log_progress("debug", "generate_report_error", f"Failed to generate report: {str(e)}")
#         raise Exception(f"Failed to generate report with Gemini: {str(e)}")
# # Helper function to save report as PDF with margins and line wrapping
# def save_report_as_pdf(report: dict, output_path: str):
#     try:
#         c = canvas.Canvas(output_path, pagesize=letter)
#         width, height = letter  # A4 size: 612 x 792 points
#         left_margin, right_margin, top_margin, bottom_margin = 72, 72, 72, 72  # 1 inch margins (72 points)
#         text_width = width - left_margin - right_margin  # Usable width for text
#         y_position = height - top_margin  # Start below top margin

#         # Helper function to add text with wrapping
#         def add_text(text, y, bold=False, indent=0):
#             nonlocal y_position
#             if y_position < bottom_margin:  # New page if below bottom margin
#                 c.showPage()
#                 y_position = height - top_margin

#             text_object = c.beginText(left_margin + indent, y_position)
#             if bold:
#                 c.setFont("Helvetica-Bold", 12)
#             else:
#                 c.setFont("Helvetica", 12)

#             # Split text into lines that fit within text_width
#             words = text.split()
#             current_line = []
#             for word in words:
#                 test_line = " ".join(current_line + [word])
#                 if c.stringWidth(test_line, "Helvetica" if not bold else "Helvetica-Bold", 12) <= text_width - indent:
#                     current_line.append(word)
#                 else:
#                     text_object.textLine(" ".join(current_line))
#                     y_position -= 15
#                     if y_position < bottom_margin:
#                         c.drawText(text_object)
#                         c.showPage()
#                         y_position = height - top_margin
#                         text_object = c.beginText(left_margin + indent, y_position)
#                         c.setFont("Helvetica" if not bold else "Helvetica-Bold", 12)
#                     current_line = [word]
#             if current_line:
#                 text_object.textLine(" ".join(current_line))
#                 y_position -= 15

#             c.drawText(text_object)

#         # Title
#         add_text("Resume Analysis Report", y_position, bold=True)
#         y_position -= 20

#         # Candidate Details
#         add_text("Candidate Details", y_position, bold=True)
#         for key, value in report["Candidate Details"].items():
#             add_text(f"- {key}: {value}", y_position, indent=10)
#         y_position -= 20

#         # Skill Match
#         add_text("Skill Match", y_position, bold=True)
#         add_text(f"Score: {report['Skill Match']['Score']}/100", y_position, indent=10)
#         add_text(f"Matches: {report['Skill Match']['Matches']}", y_position, indent=10)
#         add_text(f"Non-Matches: {report['Skill Match']['Non-Matches']}", y_position, indent=10)
#         add_text(f"Summary: {report['Skill Match']['Summary']}", y_position, indent=10)
        
#         # add_text(f"Enhancement Tips: {report['Skill Match']['Enhancement Tips']}", y_position, indent=10)
#         y_position -= 20

#         # Work Experience
#         add_text("Work Experience", y_position, bold=True)
#         add_text(f"Score: {report['Work Experience']['Score']}/100", y_position, indent=10)
#         add_text(f"Matches: {report['Work Experience']['Matches']}", y_position, indent=10)
#         add_text(f"Non-Matches: {report['Work Experience']['Non-Matches']}", y_position, indent=10)
#         add_text(f"Summary: {report['Work Experience']['Summary']}", y_position, indent=10)
#         # add_text(f"Enhancement Tips: {report['Work Experience']['Enhancement Tips']}", y_position, indent=10)
#         y_position -= 20

#         # Relevant Projects (mapped to Relevant Work Experience)
#         add_text("Relevant Projects", y_position, bold=True)
#         add_text(f"Score: {report['Relevant Projects']['Score']}/100", y_position, indent=10)
#         add_text(f"Matches: {report['Relevant Projects']['Matches']}", y_position, indent=10)
#         add_text(f"Non-Matches: {report['Relevant Projects']['Non-Matches']}", y_position, indent=10)
#         add_text(f"Summary: {report['Relevant Projects']['Summary']}", y_position, indent=10)
#         # add_text(f"Enhancement Tips: {report['Relevant Projects']['Enhancement Tips']}", y_position, indent=10)
#         y_position -= 20

#         # Education
#         add_text("Education", y_position, bold=True)
#         add_text(f"Score: {report['Education']['Score']}/100", y_position, indent=10)
#         add_text(f"Matches: {report['Education']['Matches']}", y_position, indent=10)
#         add_text(f"Non-Matches: {report['Education']['Non-Matches']}", y_position, indent=10)
#         add_text(f"Summary: {report['Education']['Summary']}", y_position, indent=10)
        
#         # add_text(f"Enhancement Tips: {report['Education']['Enhancement Tips']}", y_position, indent=10)
#         y_position -= 20

#         # Overall Summary
#         add_text("Overall Summary", y_position, bold=True)
#         add_text(report["Overall Summary"], y_position, indent=10)
#         add_text(f"Overall Score: {report['Overall Score']}/100", y_position, indent=10)

#         c.save()
#     except Exception as e:
#         raise Exception(f"Failed to save report as PDF: {str(e)}")

# # ... [rest of the unchanged functions] ...# Helper function to upload report to Supabase Storage
# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def upload_report(report_path: str, destination_path: str):
#     try:
#         with open(report_path, "rb") as f:
#             response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).upload(destination_path, f)
#             # Check response for errors
#             if hasattr(response, 'status_code') and response.status_code != 200:
#                 raise Exception(f"Upload failed with status {response.status_code}: {response.json()}")
#         return response
#     except Exception as e:
#         log_progress("upload_error", f"Failed to upload report to Supabase: {str(e)}")
#         raise Exception(f"Failed to upload report to Supabase: {str(e)}")

# # Background task to process the analysis
# def process_analysis(job_id: str, candidate_id: str, resume_path: str, job_description: str):
#     local_resume_path = None
#     local_report_path = None
#     try:
#         # Log task initiation
#         log_progress(job_id, "init", "Task started", {
#             "candidate_id": candidate_id,
#             "resume_path": resume_path,
#             "job_description": job_description
#         })

#         # Validate job_id exists in hr_jobs
#         job_response = supabase.table("hr_jobs").select("description").eq("id", job_id).execute()
#         if not job_response.data:
#             raise Exception(f"Job {job_id} not found in hr_jobs")
#         job_description = job_response.data[0]["description"]  # Use DB description if provided

#         # Download resume
#         log_progress(job_id, "download_resume", "Downloading resume from Supabase Storage")
#         local_resume_path = download_resume(resume_path)
#         log_progress(job_id, "download_resume", "Resume downloaded successfully", {
#             "local_path": local_resume_path,
#             "resume_size": os.path.getsize(local_resume_path)
#         })

#         # Extract text from resume
#         log_progress(job_id, "extract_text", "Extracting text from resume using OCR")
#         resume_text = extract_text_from_pdf(local_resume_path, job_id)
#         log_progress(job_id, "extract_text", "Text extracted successfully", {
#             "text_length": len(resume_text)
#         })

#         # Generate report
#         log_progress(job_id, "generate_report", "Generating report with Gemini")
#         report = generate_report(resume_text, job_description)
#         log_progress(job_id, "generate_report", "Report generated successfully", {
#             "overall_score": report["Overall Score"],
#             "candidate_details": report["Candidate Details"]
#         })

#         # Ensure candidate exists in hr_candidates
#         candidate_response = supabase.table("hr_candidates").select("*").eq("id", candidate_id).execute()
#         if not candidate_response.data:
#             # Safely handle missing Candidate Details
#             candidate_details = report.get("Candidate Details", {})
#             supabase.table("hr_candidates").insert({
#                 "id": candidate_id,
#                 "name": candidate_details.get("Name", "Unknown"),
#                 "email": candidate_details.get("Email", f"unknown_{candidate_id}@example.com"),
#                 "phone_number": candidate_details.get("Phone Number", None),
#                 "linkedin_url": candidate_details.get("LinkedIn", None),
#                 "github_url": candidate_details.get("GitHub", None)
#             }).execute()
#             log_progress(job_id, "create_candidate", f"Created new candidate record for {candidate_id}")

#         # Save report as PDF
#         log_progress(job_id, "save_report", "Saving report as PDF")
#         report_filename = f"report_{job_id}_{candidate_id}.pdf"
#         local_report_path = f"/tmp/{report_filename}"
#         save_report_as_pdf(report, local_report_path)
#         log_progress(job_id, "save_report", "Report saved successfully", {
#             "local_report_path": local_report_path
#         })

#         # Upload report to Supabase Storage
#         log_progress(job_id, "upload_report", "Uploading report to Supabase Storage")
#         report_destination_path = f"{SUPABASE_REPORT_PATH_PREFIX}/{job_id}/{report_filename}"
#         upload_report(local_report_path, report_destination_path)
#         log_progress(job_id, "upload_report", "Report uploaded successfully", {
#             "report_destination_path": report_destination_path
#         })

#         # Check if hr_job_candidates record exists, insert if not, update if it does
#         # Update hr_job_candidates for existing record
#         log_progress(job_id, "update_candidate", f"Updating hr_job_candidates for id: {candidate_id}")
#         existing_record = supabase.table("hr_job_candidates").select("id").eq("id", candidate_id).execute()
        
#         if not existing_record.data:
#             raise Exception(f"No hr_job_candidates record found for id: {candidate_id}")

#         # Get candidate name from report or hr_candidates
#         candidate_name = report.get("Candidate Details", {}).get("Name", None)
#         if not candidate_name:
#             # Fallback: Query hr_candidates (assuming name might still come from there)
#             candidate_response = supabase.table("hr_candidates").select("name").eq("id", candidate_id).execute()
#             candidate_name = candidate_response.data[0]["name"] if candidate_response.data else "Unknown"
        
#         candidate_data = {
#             "job_id": job_id,
#             "candidate_id": candidate_id,
#             # "resume_url": resume_path,
#             # "report_url": report_destination_path,
#             "status": "finished",
#             # "name": candidate_name,
#             "overall_score": report.get("Overall Score", 0.0),
#             "skills_score": report.get("Skill Match", {}).get("Score", 0),
#             "skills_summary": report.get("Skill Match", {}).get("Summary", ""),
#             "skills_enhancement_tips": report.get("Skill Match", {}).get("Enhancement Tips", ""),
#             "work_experience_score": report.get("Work Experience", {}).get("Score", 0),
#             "work_experience_summary": report.get("Work Experience", {}).get("Summary", ""),
#             "work_experience_enhancement_tips": report.get("Work Experience", {}).get("Enhancement Tips", ""),
#             "projects_score": report.get("Relevant Projects", {}).get("Score", 0),
#             "projects_summary": report.get("Relevant Projects", {}).get("Summary", ""),
#             "projects_enhancement_tips": report.get("Relevant Projects", {}).get("Enhancement Tips", ""),
#             "education_score": report.get("Education", {}).get("Score", 0),
#             "education_summary": report.get("Education", {}).get("Summary", ""),
#             "education_enhancement_tips": report.get("Education", {}).get("Enhancement Tips", ""),
#             "overall_summary": report.get("Overall Summary", "")
#         }

#        # Update existing record where id matches candidate_id
#         supabase.table("hr_job_candidates").update(candidate_data).eq("id", candidate_id).execute()
#         log_progress(job_id, "update_candidate", f"Updated hr_job_candidates record for id: {candidate_id}")

#         return {"status": "finished", "candidate_id": candidate_id}

#     except Exception as e:
#         log_progress(job_id, "error", f"Task failed: {str(e)}")
#         # Update status to failed if record exists
#         existing_record = supabase.table("hr_job_candidates").select("id").eq("job_id", job_id).eq("candidate_id", candidate_id).execute()
#         if existing_record.data:
#             supabase.table("hr_job_candidates").update({"status": "failed"}).eq("job_id", job_id).eq("candidate_id", candidate_id).execute()
#         return {"status": "failed", "candidate_id": candidate_id, "error": str(e)}

#     finally:
#         try:
#             log_progress(job_id, "cleanup", "Cleaning up temporary files")
#             if local_resume_path and os.path.exists(local_resume_path):
#                 os.remove(local_resume_path)
#             if local_report_path and os.path.exists(local_report_path):
#                 os.remove(local_report_path)
#             log_progress(job_id, "cleanup", "Temporary files removed successfully")
#         except Exception as e:
#             log_progress(job_id, "cleanup_error", f"Failed to clean up temporary files: {str(e)}")


# import os
# import json
# import logging
# import datetime
# from pdf2image import convert_from_path
# from pdf2image.exceptions import PDFPageCountError
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from tenacity import retry, stop_after_attempt, wait_exponential
# import pytesseract
# from docx import Document
# import textract

# # Import shared objects from config.py
# from config import gemini_model, supabase, SUPABASE_STORAGE_BUCKET, SUPABASE_REPORT_PATH_PREFIX, redis_conn

# # Set up logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # Function to log progress and store in Redis
# def log_progress(job_id: str, step: str, message: str, data: dict = None):
#     """Log a progress step and store it in Redis."""
#     log_entry = {
#         "step": step,
#         "message": message,
#         "data": data or {},
#         "timestamp": str(os.times()[4])  # System time
#     }
#     logger.info(f"Job {job_id} - {step}: {message}")
#     redis_conn.rpush(f"job_logs:{job_id}", json.dumps(log_entry))

# # Helper function to download resume from Supabase Storage
# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def download_resume(resume_path: str) -> str:
#     try:
#         response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).download(resume_path)
#         local_path = f"/tmp/{os.path.basename(resume_path)}"
#         with open(local_path, "wb") as f:
#             f.write(response)
#         return local_path
#     except Exception as e:
#         raise Exception(f"Failed to download resume from Supabase: {str(e)}")

# # Helper function to get the total number of pages in a PDF
# def get_pdf_page_count(pdf_path: str) -> int:
#     try:
#         from pdf2image import pdfinfo_from_path
#         pdf_info = pdfinfo_from_path(pdf_path)
#         return int(pdf_info["Pages"])
#     except PDFPageCountError as e:
#         raise Exception(f"Failed to determine PDF page count: {str(e)}")

# # Helper function to extract text from a file (PDF, DOCX, or DOC)
# def extract_text_from_file(file_path: str, job_id: str) -> str:
#     try:
#         file_extension = os.path.splitext(file_path)[1].lower()
#         text = ""

#         if file_extension == ".pdf":
#             total_pages = get_pdf_page_count(file_path)
#             log_progress(job_id, "extract_text", f"PDF has {total_pages} pages")
#             custom_config = r'--oem 3 --psm 6'  # OEM 3 (default), PSM 6 (assume a single uniform block of text)

#             for page_num in range(1, total_pages + 1):
#                 log_progress(job_id, "extract_text", f"Converting page {page_num}/{total_pages} to image at 200 DPI")
#                 images = convert_from_path(file_path, dpi=200, first_page=page_num, last_page=page_num)
#                 if not images:
#                     log_progress(job_id, "extract_text", f"No image generated for page {page_num}")
#                     continue

#                 page_text = pytesseract.image_to_string(images[0], config=custom_config)
#                 text += f"\n--- Page {page_num} ---\n{page_text}"
#                 log_progress(job_id, "extract_text", f"Extracted text from page {page_num}/{total_pages}", {
#                     "page_text_length": len(page_text),
#                     "page_text": page_text[:500]  # Log first 500 chars for debugging
#                 })
#                 del images

#         elif file_extension == ".docx":
#             log_progress(job_id, "extract_text", "Extracting text from DOCX")
#             doc = Document(file_path)
#             for para in doc.paragraphs:
#                 if para.text.strip():
#                     text += para.text + "\n"
#             log_progress(job_id, "extract_text", "Text extracted from DOCX successfully")

#         elif file_extension == ".doc":
#             log_progress(job_id, "extract_text", "Extracting text from DOC using textract")
#             text = textract.process(file_path).decode('utf-8')
#             log_progress(job_id, "extract_text", "Text extracted from DOC successfully")

#         else:
#             raise Exception(f"Unsupported file format: {file_extension}")

#         log_progress(job_id, "extract_text", "Text extracted successfully", {
#             "text_length": len(text),
#             "full_text": text[:1000]  # Log first 1000 chars for debugging
#         })
#         return text.strip() or "No text extracted"

#     except Exception as e:
#         raise Exception(f"Failed to extract text from file: {str(e)}")

# # Helper function to clean Gemini AI output
# def clean_gemini_output(text: str) -> str:
#     return text.strip()

# # Helper function to generate the report using Gemini API
# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def generate_report(resume_text: str, job_description: str) -> dict:
#     try:
#         # Log the input to Gemini for debugging
#         log_progress("debug", "generate_report_input", "Sending to Gemini", {
#             "resume_text": resume_text[:1000],
#             "job_description": job_description[:1000]
#         })

#         prompt = """
# Analyze this resume against the job description and return ONLY a valid JSON response with:
# - overall_match_score (percentage, 0-100)
# - matched_skills (array of objects with:
#     requirement (detailed, e.g., "Python for automation"),
#     matched ('yes', 'partial', 'no'),
#     details (specific evidence from resume or "Not mentioned" if absent))
# - summary (short plain text summary)
# - companies (array of objects with:
#     name (string, company name),
#     designation (string, role at company or "-" if not specified),
#     years (string, duration like "2019 - 2022" or "-" if not specified))
# - missing_or_weak_areas (array of strings listing gaps)
# - top_skills (array of candidate's strongest skills)
# - development_gaps (array of skills needing improvement)
# - additional_certifications (array of strings listing certifications not required by JD)
# - section_wise_scoring (array of objects with main sections, each containing:
#     {
#       section (string),
#       weightage (percentage),
#       submenus (array of { submenu (string), weightage (percentage of section), score (out of 10), weighted_score (calculated), remarks (string) })
#     })
# - candidate_name (string, extracted from resume or "Unknown" if not found)
# - email (string, extracted from resume or "" if not found)
# - phone_number (string, extracted from resume or "" if not found)
# - github (string, extracted from resume or "" if not found)
# - linkedin (string, extracted from resume or "" if not found)

# Job Description: {job_description}
# Resume: {resume_text}

# Structure section_wise_scoring with main sections and submenus:
# - Technical Skills (weightage: 40%, submenus: Core Skills 60%, Tools 40%)
# - Work Experience (weightage: 30%, submenus: Relevant Experience 70%, Duration 30%)
# - Projects (weightage: 15%, submenus: Personal Projects 50%, Professional Projects 50%)
# - Education (weightage: 10%, submenus: Degree, Certifications; weightage depends on JD:
#   - If JD requires a specific certification: Degree 30%, Certifications 70%
#   - If JD does not require a certification: Degree 50%, Certifications 50%)
# - Achievements (weightage: 5%, submenus: Awards 50%, Recognitions 50%)
# - Soft Skills (weightage: 5%, submenus: Leadership 50%, Communication 50%)

# Scoring Guidelines:
# - 'yes' (8-10/10): Clear evidence of the skill matching the JD.
# - 'partial' (4-7/10): Implied or indirect evidence.
# - 'no' (0-3/10): No evidence.
# - Infer skills from context (e.g., "Python used in automation tasks" matches "Python for automation").
# - Calculate overall_match_score as the sum of each section's weighted contribution:
#   - section_score = sum(submenu.weightage * submenu.score) / 100
#   - overall_match_score = sum(section.weightage * section_score) / 100

# For companies:
# - Extract company names, designations, and years from work experience sections.
# - If designation is not explicitly mentioned, use "-".
# - If years are not specified, use "-".
# - Example: "Senior Developer at TCS, 2019 - 2022" becomes { "name": "TCS", "designation": "Senior Developer", "years": "2019 - 2022" }

# Use symbols: ✅ for 'yes', ⚠️ for 'partial', ❌ for 'no'. Return ONLY the JSON object.
# """.format(job_description=job_description, resume_text=resume_text)

#         response = gemini_model.generate_content(prompt)
#         gemini_output = clean_gemini_output(response.text)

#         # Log Gemini's raw output for debugging
#         log_progress("debug", "generate_report_output", "Gemini response", {
#             "raw_output": gemini_output[:1000]
#         })

#         # Parse JSON response
#         try:
#             report = json.loads(gemini_output)
#         except json.JSONDecodeError as e:
#             log_progress("debug", "generate_report_error", f"Failed to parse JSON: {str(e)}")
#             raise Exception(f"Invalid JSON response from Gemini: {str(e)}")

#         # Validate required fields
#         required_fields = [
#             "overall_match_score", "matched_skills", "summary", "companies",
#             "missing_or_weak_areas", "top_skills", "development_gaps",
#             "additional_certifications", "section_wise_scoring",
#             "candidate_name", "email", "github", "linkedin"
#         ]
#         for field in required_fields:
#             if field not in report:
#                 log_progress("debug", "generate_report_error", f"Missing field in report: {field}")
#                 raise Exception(f"Missing required field in Gemini response: {field}")

#         # Normalize company names
#         def normalize_company_name(name):
#             lower_name = name.lower().strip()
#             if lower_name in ["infosys", "infosys ltd"]:
#                 return "Infosys"
#             if lower_name == "infosys infotech":
#                 return "Infosys Infotech"
#             return name.strip()

#         report["companies"] = [
#             {
#                 "name": normalize_company_name(company["name"]),
#                 "designation": company.get("designation", "-"),
#                 "years": company.get("years", "-")
#             }
#             for company in report["companies"]
#         ]

#         # Deduplicate companies by name, keeping the latest entry
#         unique_companies = []
#         seen_names = set()
#         for company in reversed(report["companies"]):
#             if company["name"] not in seen_names:
#                 unique_companies.append(company)
#                 seen_names.add(company["name"])
#         report["companies"] = list(reversed(unique_companies))

#         # Log the parsed report for debugging
#         log_progress("debug", "generate_report_parsed", "Parsed report", {
#             "report": report
#         })

#         return report

#     except Exception as e:
#         log_progress("debug", "generate_report_error", f"Failed to generate report: {str(e)}")
#         raise Exception(f"Failed to generate report with Gemini: {str(e)}")

# # Helper function to save report as PDF
# def save_report_as_pdf(report: dict, output_path: str):
#     try:
#         c = canvas.Canvas(output_path, pagesize=letter)
#         width, height = letter
#         left_margin, right_margin, top_margin, bottom_margin = 72, 72, 72, 72
#         text_width = width - left_margin - right_margin
#         y_position = height - top_margin

#         def add_text(text, y, bold=False, indent=0):
#             nonlocal y_position
#             if y_position < bottom_margin:
#                 c.showPage()
#                 y_position = height - top_margin

#             text_object = c.beginText(left_margin + indent, y_position)
#             c.setFont("Helvetica-Bold" if bold else "Helvetica", 12)

#             words = text.split()
#             current_line = []
#             for word in words:
#                 test_line = " ".join(current_line + [word])
#                 if c.stringWidth(test_line, "Helvetica" if not bold else "Helvetica-Bold", 12) <= text_width - indent:
#                     current_line.append(word)
#                 else:
#                     text_object.textLine(" ".join(current_line))
#                     y_position -= 15
#                     if y_position < bottom_margin:
#                         c.drawText(text_object)
#                         c.showPage()
#                         y_position = height - top_margin
#                         text_object = c.beginText(left_margin + indent, y_position)
#                         c.setFont("Helvetica" if not bold else "Helvetica-Bold", 12)
#                     current_line = [word]
#             if current_line:
#                 text_object.textLine(" ".join(current_line))
#                 y_position -= 15

#             c.drawText(text_object)

#         # Title
#         add_text("Resume Analysis Report", y_position, bold=True)
#         y_position -= 20

#         # Candidate Details
#         add_text("Candidate Details", y_position, bold=True)
#         add_text(f"- Name: {report.get('candidate_name', 'Unknown')}", y_position, indent=10)
#         add_text(f"- Email: {report.get('email', '')}", y_position, indent=10)
#         add_text(f"- PhoneNumber: {report.get('phone_number', '')}", y_position, indent=10)
#         add_text(f"- LinkedIn: {report.get('linkedin', '')}", y_position, indent=10)
#         add_text(f"- GitHub: {report.get('github', '')}", y_position, indent=10)
#         y_position -= 20

#         # Overall Score
#         add_text("Overall Match Score", y_position, bold=True)
#         add_text(f"{report.get('overall_match_score', 0)}/100", y_position, indent=10)
#         y_position -= 20

#         # Matched Skills
#         add_text("Matched Skills", y_position, bold=True)
#         for skill in report.get("matched_skills", []):
#             status = {"yes": "✅", "partial": "⚠️", "no": "❌"}.get(skill.get("matched", "no"))
#             add_text(f"- {skill.get('requirement', '')}: {status} {skill.get('details', '')}", y_position, indent=10)
#         y_position -= 20

#         # Summary
#         add_text("Summary", y_position, bold=True)
#         add_text(report.get("summary", ""), y_position, indent=10)
#         y_position -= 20

#         # Companies
#         add_text("Companies", y_position, bold=True)
#         for company in report.get("companies", []):
#             add_text(f"- {company.get('name', '')}: {company.get('designation', '-')}, {company.get('years', '-')}", y_position, indent=10)
#         y_position -= 20

#         # Missing or Weak Areas
#         add_text("Missing or Weak Areas", y_position, bold=True)
#         for area in report.get("missing_or_weak_areas", []):
#             add_text(f"- {area}", y_position, indent=10)
#         y_position -= 20

#         # Top Skills
#         add_text("Top Skills", y_position, bold=True)
#         for skill in report.get("top_skills", []):
#             add_text(f"- {skill}", y_position, indent=10)
#         y_position -= 20

#         # Development Gaps
#         add_text("Development Gaps", y_position, bold=True)
#         for gap in report.get("development_gaps", []):
#             add_text(f"- {gap}", y_position, indent=10)
#         y_position -= 20

#         # Additional Certifications
#         add_text("Additional Certifications", y_position, bold=True)
#         for cert in report.get("additional_certifications", []):
#             add_text(f"- {cert}", y_position, indent=10)
#         y_position -= 20

#         # Section-Wise Scoring
#         add_text("Section-Wise Scoring", y_position, bold=True)
#         for section in report.get("section_wise_scoring", []):
#             add_text(f"- {section.get('section', '')} (Weightage: {section.get('weightage', 0)}%)", y_position, indent=10)
#             for submenu in section.get("submenus", []):
#                 add_text(
#                     f"  - {submenu.get('submenu', '')}: Score {submenu.get('score', 0)}/10, "
#                     f"Remarks: {submenu.get('remarks', '')}",
#                     y_position, indent=20
#                 )
#             y_position -= 10
#         y_position -= 20

#         c.save()
#     except Exception as e:
#         raise Exception(f"Failed to save report as PDF: {str(e)}")

# # Helper function to upload report to Supabase Storage
# @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
# def upload_report(report_path: str, destination_path: str):
#     try:
#         with open(report_path, "rb") as f:
#             response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).upload(destination_path, f)
#             if hasattr(response, 'status_code') and response.status_code != 200:
#                 raise Exception(f"Upload failed with status {response.status_code}: {response.json()}")
#         return response
#     except Exception as e:
#         log_progress("upload_error", f"Failed to upload report to Supabase: {str(e)}")
#         raise Exception(f"Failed to upload report to Supabase: {str(e)}")

# # Background task to process the analysis
# def process_analysis(job_id: str, candidate_id: str, resume_path: str, job_description: str):
#     local_resume_path = None
#     local_report_path = None
#     try:
#         # Log task initiation
#         log_progress(job_id, "init", "Task started", {
#             "candidate_id": candidate_id,
#             "resume_path": resume_path,
#             "job_description": job_description
#         })

#         # Validate job_id exists in hr_jobs
#         job_response = supabase.table("hr_jobs").select("description").eq("id", job_id).execute()
#         if not job_response.data:
#             raise Exception(f"Job {job_id} not found in hr_jobs")
#         job_description = job_response.data[0]["description"]  # Use DB description if provided

#         # Download resume
#         log_progress(job_id, "download_resume", "Downloading resume from Supabase Storage")
#         local_resume_path = download_resume(resume_path)
#         log_progress(job_id, "download_resume", "Resume downloaded successfully", {
#             "local_path": local_resume_path,
#             "resume_size": os.path.getsize(local_resume_path)
#         })

#         # Extract text from resume
#         log_progress(job_id, "extract_text", "Extracting text from resume")
#         resume_text = extract_text_from_file(local_resume_path, job_id)
#         log_progress(job_id, "extract_text", "Text extracted successfully", {
#             "text_length": len(resume_text)
#         })

#         # Generate report
#         log_progress(job_id, "generate_report", "Generating report with Gemini")
#         report = generate_report(resume_text, job_description)
#         log_progress(job_id, "generate_report", "Report generated successfully", {
#             "overall_score": report["overall_match_score"],
#             "candidate_name": report["candidate_name"]
#         })

#         # Ensure candidate exists in hr_candidates
#         candidate_response = supabase.table("hr_candidates").select("*").eq("id", candidate_id).execute()
#         if not candidate_response.data:
#             supabase.table("hr_candidates").insert({
#                 "id": candidate_id,
#                 "name": report.get("candidate_name", "Unknown"),
#                 "email": report.get("email", f"unknown_{candidate_id}@example.com"),
#                 "phone_number": report.get("phone_number", None),
#                 "linkedin_url": report.get("linkedin", None),
#                 "github_url": report.get("github", None)
#             }).execute()
#             log_progress(job_id, "create_candidate", f"Created new candidate record for {candidate_id}")

#         # Save report as PDF
#         log_progress(job_id, "save_report", "Saving report as PDF")
#         report_filename = f"report_{job_id}_{candidate_id}.pdf"
#         local_report_path = f"/tmp/{report_filename}"
#         save_report_as_pdf(report, local_report_path)
#         log_progress(job_id, "save_report", "Report saved successfully", {
#             "local_report_path": local_report_path
#         })

#         # Upload report to Supabase Storage
#         log_progress(job_id, "upload_report", "Uploading report to Supabase Storage")
#         report_destination_path = f"{SUPABASE_REPORT_PATH_PREFIX}/{job_id}/{report_filename}"
#         upload_report(local_report_path, report_destination_path)
#         log_progress(job_id, "upload_report", "Report uploaded successfully", {
#             "report_destination_path": report_destination_path
#         })

#         # Step 1: Prepare resume_analysis payload
#         resume_payload = {
#             "job_id": job_id,
#             "candidate_id": candidate_id,
#             "resume_text": resume_text or None,
#             "overall_score": round(report.get("overall_match_score", 0)),
#             "matched_skills": report.get("matched_skills", []),
#             "summary": report.get("summary", None),
#             "missing_or_weak_areas": report.get("missing_or_weak_areas", []),
#             "top_skills": report.get("top_skills", []),
#             "development_gaps": report.get("development_gaps", []),
#             "additional_certifications": report.get("additional_certifications", []),
#             "section_wise_scoring": report.get("section_wise_scoring", {}),
#             "candidate_name": report.get("candidate_name", "Unknown"),
#             "email": report.get("email", ""),
#             "phone_number": report.get("phone_number",""),
#             "github": report.get("github", ""),
#             "linkedin": report.get("linkedin", ""),
#             "updated_at": datetime.datetime.utcnow().isoformat()
#         }

#         # Step 2: Upsert resume_analysis
#         log_progress(job_id, "save_resume_analysis", "Upserting resume_analysis")
#         response = supabase.table("resume_analysis").upsert(
#             resume_payload,
#             options={"on_conflict": ["job_id", "candidate_id"]}
#         ).execute()
#         if not response.data:
#             raise Exception("Failed to upsert resume_analysis")

#         # Step 3: Save and check company associations in candidate_companies
#         company_entries = []
#         for company in report.get("companies", []):
#             # Check if company exists
#             company_response = supabase.table("companies").select("id").eq("name", company["name"]).execute()
#             if company_response.data:
#                 company_id = company_response.data[0]["id"]
#             else:
#                 # Insert new company
#                 new_company = supabase.table("companies").insert({"name": company["name"]}).execute()
#                 if not new_company.data:
#                     raise Exception(f"Failed to insert company: {company['name']}")
#                 company_id = new_company.data[0]["id"]

#             company_entries.append({
#                 "candidate_id": candidate_id,
#                 "job_id": job_id,
#                 "company_id": company_id,
#                 "designation": company.get("designation", "-"),
#                 "years": company.get("years", "-")
#             })

#         if company_entries:
#             log_progress(job_id, "save_candidate_companies", "Upserting candidate_companies")
#             company_response = supabase.table("candidate_companies").upsert(
#                 company_entries,
#                 options={"on_conflict": ["candidate_id", "job_id", "company_id"]}
#             ).execute()
#             if not company_response.data:
#                 raise Exception("Failed to upsert candidate_companies")

#         return {"status": "finished", "candidate_id": candidate_id}

#     except Exception as e:
#         log_progress(job_id, "error", f"Task failed: {str(e)}")
#         # Update status to failed in resume_analysis if record exists
#         existing_record = supabase.table("resume_analysis").select("candidate_id").eq("job_id", job_id).eq("candidate_id", candidate_id).execute()
#         if existing_record.data:
#             supabase.table("resume_analysis").update({"status": "failed"}).eq("job_id", job_id).eq("candidate_id", candidate_id).execute()
#         return {"status": "failed", "candidate_id": candidate_id, "error": str(e)}

#     finally:
#         try:
#             log_progress(job_id, "cleanup", "Cleaning up temporary files")
#             if local_resume_path and os.path.exists(local_resume_path):
#                 os.remove(local_resume_path)
#             if local_report_path and os.path.exists(local_report_path):
#                 os.remove(local_report_path)
#             log_progress(job_id, "cleanup", "Temporary files removed successfully")
#         except Exception as e:
#             log_progress(job_id, "cleanup_error", f"Failed to clean up temporary files: {str(e)}")


from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
import os
import json
import re
import logging
import datetime
from pdf2image import convert_from_path
from pdf2image.exceptions import PDFPageCountError
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from tenacity import retry, stop_after_attempt, wait_exponential
import pytesseract
from docx import Document
import textract
import ast
import traceback

# Import shared objects from config.py
from config import gemini_model, supabase, SUPABASE_STORAGE_BUCKET, SUPABASE_REPORT_PATH_PREFIX, redis_conn

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Define Styles (Can be done outside the function for reuse) ---
styles = getSampleStyleSheet()

# Custom Styles - Refined for Appeal
styles.add(ParagraphStyle(name='ReportTitle',
                          parent=styles['h1'],
                          alignment=TA_CENTER,
                          fontSize=20,
                          spaceBottom=24,
                          textColor=colors.HexColor('#3A0CA3'))) # Deep Purple Title

styles.add(ParagraphStyle(name='SectionHeading',
                          parent=styles['h2'],
                          fontSize=14,
                          spaceBefore=16,
                          spaceAfter=8,
                          textColor=colors.HexColor('#480CA8'), # Royal Blue Heading
                          alignment=TA_LEFT))

styles.add(ParagraphStyle(name='SubHeading',
                          parent=styles['h3'],
                          fontSize=11,
                          spaceBefore=10,
                          spaceAfter=5,
                          textColor=colors.HexColor('#480CA8'), # Another Purple Shade
                          fontName='Helvetica-Bold',
                          alignment=TA_LEFT))

styles.add(ParagraphStyle(name='Body',
                          parent=styles['Normal'],
                          alignment=TA_LEFT,
                          fontSize=10,
                          leading=14, # Line spacing
                          textColor=colors.darkslategray))

styles.add(ParagraphStyle(name='ListItem',
                          parent=styles['Body'],
                          leftIndent=20,
                          spaceBefore=2,
                          spaceAfter=2)) # For bullet points

styles.add(ParagraphStyle(name='ScoreHighlight',
                          parent=styles['Normal'],
                          alignment=TA_RIGHT,
                          fontSize=22, # Larger score
                          fontName='Helvetica-Bold',
                          textColor=colors.HexColor('#3A0CA3'))) # Match title color

styles.add(ParagraphStyle(name='ScoreLabel',
                          parent=styles['Normal'],
                          alignment=TA_LEFT,
                          fontSize=12,
                          fontName='Helvetica-Bold',
                          textColor=colors.HexColor('#4361EE'))) # Match heading color

styles.add(ParagraphStyle(name='TableHeader',
                           parent=styles['Normal'],
                           fontName='Helvetica-Bold',
                           fontSize=9,
                           alignment=TA_LEFT,
                           textColor=colors.white))

styles.add(ParagraphStyle(name='TableCell',
                           parent=styles['Normal'],
                           fontSize=9,
                           leading=11))

styles.add(ParagraphStyle(name='WatermarkStyle', # Can be reused if needed elsewhere
                          parent=styles['Normal'],
                          alignment=TA_RIGHT,
                          fontSize=8,
                          textColor=colors.Color(0,0,0, alpha=0.15))) # Faint watermark

# --- Watermark Function ---
def add_watermark(canvas, doc):
    """Adds 'hrumbles.ai' watermark to each page."""
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.Color(0,0,0, alpha=0.15)) # Faint color
    # Position in top-right corner
    canvas.drawRightString(doc.pagesize[0] - 0.5*inch, doc.pagesize[1] - 0.5*inch, "hrumbles.ai")
    canvas.restoreState()


# Function to log progress and store in Redis
def log_progress(job_id: str, step: str, message: str, data: dict = None):
    """Log a progress step and store in Redis."""
    log_entry = {
        "step": step,
        "message": message,
        "data": data or {},
        "timestamp": str(os.times()[4])  # System time
    }
    logger.info(f"Job {job_id} - {step}: {message}")
    redis_conn.rpush(f"job_logs:{job_id}", json.dumps(log_entry))

# Helper function to download resume from Supabase Storage
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def download_resume(resume_path: str) -> str:
    try:
        response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).download(resume_path)
        local_path = f"/tmp/{os.path.basename(resume_path)}"
        with open(local_path, "wb") as f:
            f.write(response)
        return local_path
    except Exception as e:
        raise Exception(f"Failed to download resume from Supabase: {str(e)}")

# Helper function to get the total number of pages in a PDF
def get_pdf_page_count(pdf_path: str) -> int:
    try:
        from pdf2image import pdfinfo_from_path
        pdf_info = pdfinfo_from_path(pdf_path)
        return int(pdf_info["Pages"])
    except PDFPageCountError as e:
        raise Exception(f"Failed to determine PDF page count: {str(e)}")

# Helper function to extract text from a file (PDF, DOCX, or DOC)
def extract_text_from_file(file_path: str, job_id: str) -> str:
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        text = ""

        if file_extension == ".pdf":
            total_pages = get_pdf_page_count(file_path)
            log_progress(job_id, "extract_text", f"PDF has {total_pages} pages")
            custom_config = r'--oem 3 --psm 6'  # OEM 3 (default), PSM 6 (assume a single uniform block of text)

            for page_num in range(1, total_pages + 1):
                log_progress(job_id, "extract_text", f"Converting page {page_num}/{total_pages} to image at 200 DPI")
                images = convert_from_path(file_path, dpi=200, first_page=page_num, last_page=page_num)
                if not images:
                    log_progress(job_id, "extract_text", f"No image generated for page {page_num}")
                    continue

                page_text = pytesseract.image_to_string(images[0], config=custom_config)
                text += f"\n--- Page {page_num} ---\n{page_text}"
                log_progress(job_id, "extract_text", f"Extracted text from page {page_num}/{total_pages}", {
                    "page_text_length": len(page_text),
                    "page_text": page_text[:500]  # Log first 500 chars for debugging
                })
                del images

        elif file_extension == ".docx":
            log_progress(job_id, "extract_text", "Extracting text from DOCX")
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            log_progress(job_id, "extract_text", "Text extracted from DOCX successfully")

        elif file_extension == ".doc":
            log_progress(job_id, "extract_text", "Extracting text from DOC using textract")
            text = textract.process(file_path).decode('utf-8')
            log_progress(job_id, "extract_text", "Text extracted from DOC successfully")

        else:
            raise Exception(f"Unsupported file format: {file_extension}")

        log_progress(job_id, "extract_text", "Text extracted successfully", {
            "text_length": len(text),
            "full_text": text[:1000]  # Log first 1000 chars for debugging
        })
        return text.strip() or "No text extracted"

    except Exception as e:
        raise Exception(f"Failed to extract text from file: {str(e)}")



# --- Keep the SIMPLIFIED clean_gemini_output ---
def clean_gemini_output(text: str) -> str:
    text = text.strip()
    # Remove code fences (common issue)
    if text.startswith('```json'):
        text = text[7:].rstrip('```')
    elif text.startswith('```'):
        text = text[3:].rstrip('```')
    text = text.strip()

    # Extract JSON object between the first { and the last }
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        text = text[start:end+1]
    else:
         log_progress("debug", "clean_gemini_output_warning", "No JSON object markers ({}) found in text", {"text_preview": text[:200]})
         # Let json.loads handle potential errors if no markers found

    # Remove common control characters (safer)
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    # Remove trailing commas before } or ] (relatively safe)
    text = re.sub(r',\s*([}\]])', r'\1', text)
    return text

# Helper function to generate the report using Gemini API
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def generate_report(resume_text: str, job_description: str, job_id: str) -> dict: # Added job_id parameter
    # Use a unique identifier for this specific attempt for easier log tracing
    attempt_id = os.urandom(4).hex()
    current_step = "init"
    try:
        current_step = "log_start"
        # DETAILED LOG: Function entry
        log_progress(job_id, f"generate_report_start_{attempt_id}", "Entering generate_report function")

        current_step = "log_raw_inputs"
        # DETAILED LOG: Log raw inputs before escaping
        log_progress(job_id, f"generate_report_raw_inputs_{attempt_id}", "Raw input previews", {
            "resume_preview": resume_text[:200],
            "jd_preview": job_description[:200]
        })

        # --- Your existing brace escaping ---
        current_step = "escape_inputs"
        escaped_resume = resume_text.replace('{', '{{').replace('}', '}}')
        escaped_job_desc = job_description.replace('{', '{{').replace('}', '}}')
        # DETAILED LOG: Log escaped inputs
        log_progress(job_id, f"generate_report_escaped_{attempt_id}", "Inputs escaped for formatting", {
            "resume_preview": escaped_resume[:200],
            "jd_preview": escaped_job_desc[:200]
        })

        # --- Your existing input log (now with attempt_id) ---
        current_step = "log_input_preview"
        log_progress(job_id, f"generate_report_input_{attempt_id}", "Preparing prompt", { # Added attempt_id
            "resume_text_preview": escaped_resume[:1000],
            "job_description_preview": escaped_job_desc[:1000]
        })

        # --- Your existing prompt definition ---
        prompt_template = """
Analyze this resume against the job description and return ONLY a valid JSON response with:
- overall_match_score (percentage, 0-100)
- matched_skills (array of objects with:
    requirement (detailed, e.g., "Python for automation"),
    matched ('yes', 'partial', 'no'),
    details (specific evidence from resume or "Not mentioned" if absent))
- summary (short plain text summary)
- companies (array of objects with:
    name (string, company name),
    designation (string, role at company or "-" if not specified),
    years (string, duration like "2019 - 2022" or "-" if not specified))
- missing_or_weak_areas (array of strings listing gaps)
- top_skills (array of candidate's strongest skills)
- development_gaps (array of skills needing improvement)
- additional_certifications (array of strings listing certifications not required by JD)
- section_wise_scoring (array of objects with main sections, each containing:
    {{
      section (string),
      weightage (percentage),
      submenus (array of {{ submenu (string), weightage (percentage of section), score (out of 10), weighted_score (calculated), remarks (string) }})
    }})
- candidate_name (string, extracted from resume or "Unknown" if not found)
- email (string, extracted from resume or "" if not found)
- phone_number (string, extracted from resume or "" if not found)
- github (string, extracted from resume or "" if not found)
- linkedin (string, extracted from resume or "" if not found)

Job Description: {job_description}
Resume: {resume_text}

Structure section_wise_scoring with main sections and submenus:
- Technical Skills (weightage: 40%, submenus: Core Skills 60%, Tools 40%)
- Work Experience (weightage: 30%, submenus: Relevant Experience 70%, Duration 30%)
- Projects (weightage: 15%, submenus: Personal Projects 50%, Professional Projects 50%)
- Education (weightage: 10%, submenus: Degree, Certifications; weightage depends on JD:
  - If JD requires a specific certification: Degree 30%, Certifications 70%
  - If JD does not require a certification: Degree 50%, Certifications 50%)
- Achievements (weightage: 5%, submenus: Awards 50%, Recognitions 50%)
- Soft Skills (weightage: 5%, submenus: Leadership 50%, Communication 50%)

Scoring Guidelines:
- 'yes' (8-10/10): Clear evidence of the skill matching the JD.
- 'partial' (4-7/10): Implied or indirect evidence.
- 'no' (0-3/10): No evidence.
- Infer skills from context (e.g., "Python used in automation tasks" matches "Python for automation").
- Calculate overall_match_score as the sum of each section's weighted contribution:
  - section_score = sum(submenu.weightage * submenu.score) / 100
  - overall_match_score = sum(section.weightage * section_score) / 100

For companies:
- Extract company names, designations, and years from work experience sections.
- If designation is not explicitly mentioned, use "-".
- If years are not specified, use "-".
- Example: "Senior Developer at TCS, 2019 - 2022" becomes {{ "name": "TCS", "designation": "Senior Developer", "years": "2019 - 2022" }}

Use symbols: ✅ for 'yes', ⚠️ for 'partial', ❌ for 'no'. IMPORTANT: Ensure the output is ONLY a single, valid JSON object. All string values within the JSON must be properly escaped according to JSON standards (e.g., use \\" for quotes inside strings, \\\\ for backslashes, etc.). Do NOT include any explanatory text before or after the JSON object.
"""
        current_step = "before_format"
        # DETAILED LOG: Before formatting prompt
        log_progress(job_id, f"generate_report_before_format_{attempt_id}", "About to format the prompt string")

        # --- Your existing prompt formatting ---
        prompt = prompt_template.format(job_description=escaped_job_desc, resume_text=escaped_resume)

        current_step = "after_format"
        # DETAILED LOG: After formatting prompt
        log_progress(job_id, f"generate_report_after_format_{attempt_id}", "Prompt string formatted", {"prompt_preview": prompt[:500]})

        # --- Your existing Gemini call ---
        current_step = "before_api_call"
        # DETAILED LOG: Before calling Gemini API
        log_progress(job_id, f"generate_report_before_api_call_{attempt_id}", "About to call gemini_model.generate_content()")
        response = gemini_model.generate_content(prompt)
        current_step = "after_api_call"
        # DETAILED LOG: After calling Gemini API
        log_progress(job_id, f"generate_report_after_api_call_{attempt_id}", "Returned from gemini_model.generate_content()")

        raw_gemini_text = response.text # Get the raw text

        # --- Your existing raw output log (now with attempt_id) ---
        current_step = "log_raw_output"
        log_progress(job_id, f"generate_report_raw_output_{attempt_id}", "Raw Gemini response received", { # Added attempt_id
            "raw_output_preview": raw_gemini_text[:2000]
        })

        # --- Your existing cleaning call ---
        current_step = "clean_output"
        gemini_output = clean_gemini_output(raw_gemini_text)
        # DETAILED LOG: After cleaning output
        log_progress(job_id, f"generate_report_cleaned_output_{attempt_id}", "Cleaned Gemini response", { # Added attempt_id
            "cleaned_output_preview": gemini_output[:2000]
        })

        # --- Your existing JSON parsing block ---
        current_step = "parse_json"
        try:
            report = json.loads(gemini_output)
        except json.JSONDecodeError as e:
            # DETAILED LOG: JSON parsing failed
            log_progress(job_id, f"generate_report_parse_error_{attempt_id}", f"Failed to parse JSON: {str(e)}", { # Added attempt_id
                "parsing_error": str(e),
                "cleaned_output_preview": gemini_output[:2000]
            })
            current_step = "parse_fallback"
            # DETAILED LOG: Attempting fallback parsing
            log_progress(job_id, f"generate_report_fallback_{attempt_id}", "Attempting fallback parsing with ast.literal_eval") # Added attempt_id
            try:
                report = ast.literal_eval(gemini_output)
                if not isinstance(report, dict):
                    raise ValueError("Fallback parsing did not yield a dictionary")
                # DETAILED LOG: Fallback parsing succeeded
                log_progress(job_id, f"generate_report_fallback_success_{attempt_id}", "Fallback parsing successful") # Added attempt_id
            except (ValueError, SyntaxError) as fallback_e:
                # DETAILED LOG: Fallback parsing failed
                log_progress(job_id, f"generate_report_fallback_error_{attempt_id}", f"Fallback parsing failed: {str(fallback_e)}", { # Added attempt_id
                    "fallback_error": str(fallback_e),
                    "cleaned_output_preview": gemini_output[:2000]
                })
                raise Exception(f"Invalid JSON response from Gemini after cleaning and fallback: Original error: {str(e)}. Fallback error: {str(fallback_e)}") from e # Keep combined error

        # --- Your existing parsed report log (now with attempt_id) ---
        current_step = "log_parsed_report"
        log_progress(job_id, f"generate_report_parsed_{attempt_id}", "Parsed report structure", { # Added attempt_id
            "report_keys": list(report.keys())
        })

        # --- Your existing validation ---
        current_step = "validate_report"
        required_fields = [
            "overall_match_score", "matched_skills", "summary", "companies",
            "missing_or_weak_areas", "top_skills", "development_gaps",
            "additional_certifications", "section_wise_scoring",
            "candidate_name", "email","phone_number", "github", "linkedin"
        ]
        for field in required_fields:
            if field not in report:
                # DETAILED LOG: Validation failed
                log_progress(job_id, f"generate_report_validation_error_{attempt_id}", f"Missing field in report: {field}", { # Added attempt_id
                    "report_keys": list(report.keys())
                })
                raise Exception(f"Missing required field in Gemini response: {field}")

        # --- Your existing company normalization ---
        current_step = "normalize_companies"
        # Define normalization func locally or ensure it's imported/available
        # Using a local definition here to be self-contained:
        def normalize_company_name_local(name):
            if not isinstance(name, str): # Handle potential non-string names
                return "-"
            lower_name = name.lower().strip()
            # Add your specific rules here if needed
            if lower_name in ["infosys", "infosys ltd"]:
                return "Infosys"
            if lower_name == "infosys infotech":
                return "Infosys Infotech"
            # General cleanup (optional, adjust as needed)
            normalized = re.sub(r'\s*(ltd|limited|inc|corp|corporation|llc|co)\.?\s*$', '', lower_name, flags=re.IGNORECASE)
            normalized = re.sub(r'[^\w\s-]', '', normalized) # Keep hyphens maybe?
            normalized = ' '.join(normalized.split()) # Consolidate whitespace
            return normalized.strip() if normalized else name.strip() # Return original if normalization results in empty

        # Ensure report["companies"] exists and is a list before processing
        if "companies" in report and isinstance(report["companies"], list):
            processed_companies = []
            for company in report["companies"]:
                 # Ensure company is a dict and has 'name' before processing
                if isinstance(company, dict) and "name" in company:
                    processed_companies.append({
                        "name": company["name"], # Keep original name in report for display?
                        "normalized_name_for_dedup": normalize_company_name_local(company["name"]), # Use a temp key for dedup logic
                        "designation": company.get("designation", "-"),
                        "years": company.get("years", "-")
                    })
                # else: log warning maybe? Skipping invalid company entry

            # Deduplicate based on the temporary normalized name, keeping the latest entry
            unique_companies_final = []
            seen_normalized_names = set()
            for company in reversed(processed_companies):
                norm_name = company["normalized_name_for_dedup"]
                if norm_name not in seen_normalized_names:
                    # Remove the temporary key before adding to final list
                    del company["normalized_name_for_dedup"]
                    unique_companies_final.append(company)
                    seen_normalized_names.add(norm_name)
            report["companies"] = list(reversed(unique_companies_final))
        else:
             # Handle case where 'companies' is missing or not a list
             report["companies"] = []
             log_progress(job_id, f"generate_report_normalization_warning_{attempt_id}", "Key 'companies' missing or not a list in Gemini report", {
                 "report_keys": list(report.keys())
             })


        # --- Your existing final parsed report log (now with attempt_id) ---
        # This might log large amounts of data, consider limiting it
        current_step = "log_final_report"
        try:
            report_preview = json.dumps(report)[:2000] # Log preview
        except Exception:
            report_preview = "Error creating report preview"
        log_progress(job_id, f"generate_report_final_parsed_{attempt_id}", "Final processed report preview", { # Added attempt_id
            "report_preview": report_preview
        })

        current_step = "final_success"
        # DETAILED LOG: Final success
        log_progress(job_id, f"generate_report_success_{attempt_id}", "Report generated and parsed successfully")
        return report

    # --- Updated Exception Block ---
    except Exception as e:
        # Log the exception with traceback and the last known step
        tb_str = traceback.format_exc()
        # Use logger directly for critical error logging to avoid log_progress issues
        logger.error(f"Job {job_id} - generate_report_error_{attempt_id}: Exception caught in generate_report at step '{current_step}'")
        logger.error(f"Job {job_id} - Error Type: {type(e).__name__}")
        logger.error(f"Job {job_id} - Error Message: {str(e)}")
        logger.error(f"Job {job_id} - Traceback:\n{tb_str}")

        # Also try to log via log_progress for Redis record
        log_progress(job_id, f"generate_report_error_{attempt_id}", f"Exception caught at step '{current_step}': {str(e)}", {
            "error_type": type(e).__name__,
            "error_message": str(e),
            "last_step": current_step,
            "traceback": tb_str
        })
        # Re-raise the exception for the retry mechanism
        raise Exception(f"Failed to generate report at step '{current_step}': {str(e)}") from e

# --- UPDATED save_report_as_pdf using Platypus for Visual Appeal ---
def save_report_as_pdf(report: dict, output_path: str, job_id: str): # Added job_id for logging
    try:
        log_progress(job_id, "pdf_generation_start", "Starting PDF generation with Platypus")
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                                leftMargin=0.75*inch, rightMargin=0.75*inch,
                                topMargin=1.0*inch, bottomMargin=0.75*inch)
        story = []
        bullet = '•' # Bullet character

        # --- Build Story (Content) ---

        # 1. Report Title
        story.append(Paragraph("Resume Analysis Report", styles['ReportTitle']))

        # 2. Candidate Details Table
        story.append(Paragraph("Candidate Details", styles['SectionHeading']))
        candidate_data = [
            [Paragraph("<b>Name:</b>", styles['TableCell']), Paragraph(report.get('candidate_name', 'N/A'), styles['TableCell'])],
            [Paragraph("<b>Email:</b>", styles['TableCell']), Paragraph(report.get('email', 'N/A'), styles['TableCell'])],
            [Paragraph("<b>Phone:</b>", styles['TableCell']), Paragraph(report.get('phone_number', 'N/A'), styles['TableCell'])],
            [Paragraph("<b>LinkedIn:</b>", styles['TableCell']), Paragraph(report.get('linkedin', 'N/A'), styles['TableCell'])],
            [Paragraph("<b>GitHub:</b>", styles['TableCell']), Paragraph(report.get('github', 'N/A'), styles['TableCell'])],
        ]
        candidate_table = Table(candidate_data, colWidths=[1.2*inch, 5.8*inch])
        candidate_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(candidate_table)
        story.append(Spacer(1, 0.25*inch))

        # 3. Overall Score (Highlighted Box)
        score_data = [[
            Paragraph("Overall Match Score", styles['ScoreLabel']),
            Paragraph(f"{report.get('overall_match_score', 0)}%", styles['ScoreHighlight'])
        ]]
        score_table = Table(score_data, colWidths=[5*inch, 2*inch])
        score_table.setStyle(TableStyle([
             ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
             ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F0F3FF')), # Light Blue/Lavender background
             ('LEFTPADDING', (0, 0), (0, 0), 12),
             ('RIGHTPADDING', (1, 0), (1, 0), 12),
             ('TOPPADDING', (0, 0), (-1, -1), 10),
             ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
             # ('ROUNDEDCORNERS', (0, 0), (-1,-1), 6), # Requires reportlab >= 3.6, uncomment if available
             ('LINEBELOW', (0,0), (-1,0), 1, colors.HexColor('#4361EE')), # Underline below score box
        ]))
        story.append(score_table)
        story.append(Spacer(1, 0.25*inch))

        # 4. Overall Summary
        story.append(Paragraph("Overall Summary", styles['SectionHeading']))
        story.append(Paragraph(report.get('summary', "N/A"), styles['Body']))
        story.append(Spacer(1, 0.25*inch))

        # 5. Skills Analysis (Side-by-side Table)
        story.append(Paragraph("Skills Overview", styles['SectionHeading']))

        top_skills_list = report.get("top_skills", [])
        missed_skills_list = report.get("missing_or_weak_areas", [])

        # Create Paragraph lists for each cell
        top_skills_flowables = [Paragraph("<b>Top Skills</b>", styles['SubHeading'])]
        if top_skills_list:
            for skill in top_skills_list:
                top_skills_flowables.append(Paragraph(f"{bullet} {skill}", styles['ListItem']))
        else:
            top_skills_flowables.append(Paragraph("N/A", styles['Body']))

        missed_skills_flowables = [Paragraph("<b>Missed / Weak Areas</b>", styles['SubHeading'])]
        if missed_skills_list:
            for area in missed_skills_list:
                missed_skills_flowables.append(Paragraph(f"{bullet} {area}", styles['ListItem']))
        else:
             missed_skills_flowables.append(Paragraph("N/A", styles['Body']))

        # Add KeepTogether to try and keep skill lists from breaking across pages awkwardly
        skills_data = [[top_skills_flowables, missed_skills_flowables]]
        
        skills_table = Table(skills_data, colWidths=[3.5*inch, 3.5*inch])
        skills_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#D9E2FF')), # Light blue border
            # ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.lightgrey), # Optional inner line
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(skills_table)
        story.append(Spacer(1, 0.25*inch))


        # 6. Matched Skills (Detailed Table)
        story.append(Paragraph("Detailed Skill Match", styles['SectionHeading']))
        matched_skills_list = report.get("matched_skills", [])
        if matched_skills_list:
            matched_skills_data = [[
                Paragraph("Requirement", styles['TableHeader']),
                Paragraph("Match", styles['TableHeader']),
                Paragraph("Evidence / Details", styles['TableHeader'])
            ]]
            status_map = {"yes": "✅ Yes", "partial": "⚠️ Partial", "no": "❌ No"}
            for skill in matched_skills_list:
                status = status_map.get(skill.get("matched", "no"), "❓")
                req = Paragraph(skill.get('requirement', 'N/A'), styles['TableCell'])
                det = Paragraph(skill.get('details', 'N/A'), styles['TableCell'])
                matched_skills_data.append([req, Paragraph(status, styles['TableCell']), det])

            matched_table = Table(matched_skills_data, colWidths=[2.5*inch, 0.8*inch, 3.7*inch])
            matched_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4361EE')), # Header background
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 6),
                # Alternating row colors
                # ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F3FF')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 4), # Padding for data rows
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
            ]))
            story.append(matched_table)
        else:
            story.append(Paragraph("No detailed skill matching data available.", styles['Body']))
        story.append(Spacer(1, 0.25*inch))


        # 7. Companies Mentioned
        story.append(Paragraph("Companies Mentioned", styles['SectionHeading']))
        companies_list = report.get("companies", [])
        if companies_list:
             company_items = []
             for company in companies_list:
                 name = company.get('name', 'N/A')
                 des = company.get('designation', '-')
                 yrs = company.get('years', '-')
                 company_items.append(Paragraph(f"{bullet} {name} ({des}, {yrs})", styles['ListItem']))
             story.extend(company_items) # Add list items to story
        else:
             story.append(Paragraph("N/A", styles['Body']))
        story.append(Spacer(1, 0.25*inch))

        # 8. Development Gaps
        story.append(Paragraph("Development Gaps", styles['SectionHeading']))
        gaps_list = report.get("development_gaps", [])
        if gaps_list:
            gap_items = []
            for gap in gaps_list:
                 gap_items.append(Paragraph(f"{bullet} {gap}", styles['ListItem']))
            story.extend(gap_items)
        else:
            story.append(Paragraph("N/A", styles['Body']))
        story.append(Spacer(1, 0.25*inch))

        # 9. Additional Certifications
        story.append(Paragraph("Additional Certifications", styles['SectionHeading']))
        certs_list = report.get("additional_certifications", [])
        if certs_list:
            cert_items = []
            for cert in certs_list:
                 cert_items.append(Paragraph(f"{bullet} {cert}", styles['ListItem']))
            story.extend(cert_items)
        else:
            story.append(Paragraph("N/A", styles['Body']))
        story.append(Spacer(1, 0.25*inch))


        # 10. Section-Wise Scoring (Improved Table Layout)
        story.append(Paragraph("Section-Wise Scoring", styles['SectionHeading']))
        scoring_list = report.get("section_wise_scoring", [])
        if scoring_list:
            scoring_data = [[
                Paragraph("Section", styles['TableHeader']),
                Paragraph("Sub-Section", styles['TableHeader']),
                Paragraph("Score", styles['TableHeader']),
                Paragraph("Remarks", styles['TableHeader']),
            ]]
            for section in scoring_list:
                sec_name = section.get('section', 'N/A')
                sec_weight = section.get('weightage', 0)
                first_row_for_section = True
                for submenu in section.get("submenus", []):
                    sub_name = submenu.get('submenu', 'N/A')
                    sub_score = submenu.get('score', 0)
                    sub_rem = submenu.get('remarks', '')
                    # Display main section only on the first row of its submenus
                    section_display = Paragraph(f"<b>{sec_name}</b><br/>({sec_weight}%)", styles['TableCell']) if first_row_for_section else ""
                    scoring_data.append([
                         section_display,
                         Paragraph(sub_name, styles['TableCell']),
                         Paragraph(f"{sub_score}/10", styles['TableCell']),
                         Paragraph(sub_rem, styles['TableCell'])
                    ])
                    first_row_for_section = False # Only show main section name once

            scoring_table = Table(scoring_data, colWidths=[1.5*inch, 1.7*inch, 0.8*inch, 3.0*inch])
            scoring_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4361EE')), # Header background
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 6),
                # ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F3FF')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 4), # Padding for data rows
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                # Span the Section cell across rows if needed (more complex)
            ]))
            story.append(scoring_table)
        else:
            story.append(Paragraph("N/A", styles['Body']))
        story.append(Spacer(1, 0.25*inch))


        # --- Build the PDF ---
        log_progress(job_id, "pdf_generation_build", "Building PDF document with Platypus")
        doc.build(story, onFirstPage=add_watermark, onLaterPages=add_watermark)
        log_progress(job_id, "pdf_generation_success", "PDF generated successfully")

    except Exception as e:
        tb_str = traceback.format_exc()
        logger.error(f"Job {job_id} - Failed to save report as PDF: {str(e)}")
        logger.error(f"Job {job_id} - Traceback:\n{tb_str}")
        log_progress(job_id, "pdf_generation_error", f"Failed to save report as PDF: {str(e)}", {
            "error_type": type(e).__name__,
            "error_message": str(e),
            "traceback": tb_str
        })
        raise Exception(f"Failed to save report as PDF: {str(e)}") from e

# END OF NEW save_report_as_pdf FUNCTION
# Helper function to upload report to Supabase Storage
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def upload_report(report_path: str, destination_path: str, job_id: str):
    try:
        # --- NEW LOGGING ---
        file_exists = os.path.exists(report_path)
        file_size = os.path.getsize(report_path) if file_exists else -1
        log_progress(job_id, "upload_report_info", "Preparing to upload report", {
            "local_report_path": report_path,
            "destination_path": destination_path,
            "local_file_exists": file_exists,
            "local_file_size_bytes": file_size
        })
        if not file_exists or file_size <= 0:
             log_progress(job_id, "upload_report_error", "Local report file missing or empty", {"path": report_path, "size": file_size})
             raise Exception(f"Local report file missing or empty: {report_path}")
        # --- END NEW LOGGING ---
        with open(report_path, "rb") as f:
            response = supabase.storage.from_(SUPABASE_STORAGE_BUCKET).upload(
                path=destination_path,
                file=f,
                # --- Ensure file_options are included ---
                file_options={
                    "upsert": "true", # Allow overwriting
                    "content-type": "application/pdf" # Set correct MIME type
                }
                # ---
            )
            log_progress(job_id, "upload_report_success", "Upload API call successful", {"destination": destination_path})
            if hasattr(response, 'status_code') and response.status_code != 200:
                raise Exception(f"Upload failed with status {response.status_code}: {response.json()}")
        return response     
    except Exception as e:
        tb_str = traceback.format_exc()
        log_progress(job_id, "upload_report_exception", f"Failed to upload report to Supabase: {str(e)}", {
             "local_report_path": report_path,
             "destination_path": destination_path,
             "error_type": type(e).__name__,
             "error_message": str(e),
             "traceback": tb_str
        })
        raise Exception(f"Failed to upload report to Supabase ({type(e).__name__}): {str(e)}") from e

# Helper function to normalize company names
def normalize_company_name(name: str) -> str:
    """Normalize company name by converting to lowercase and removing common suffixes."""
    if not name:
        return ""
    # Convert to lowercase and strip whitespace
    normalized = name.lower().strip()
    # Remove common suffixes (ltd, limited, inc, etc.)
    normalized = re.sub(r'\s*(ltd|limited|inc|corp|corporation|llc|co)\.?\s*$', '', normalized, flags=re.IGNORECASE)
    # Remove punctuation
    normalized = re.sub(r'[^\w\s]', '', normalized)
    return normalized
# Background task to process the analysis
# PASTE THIS ENTIRE FUNCTION INTO tasks.py, REPLACING YOUR CURRENT process_analysis

# PASTE THIS ENTIRE FUNCTION INTO tasks.py, REPLACING THE OLD process_analysis

# Background task to process the analysis
def process_analysis(job_id: str, candidate_id: str, resume_path: str, job_description_from_request: str): # Renamed input JD
    logger.info("Starting process_analysis for job_id: %s, candidate_id: %s", job_id, candidate_id)
    local_resume_path = None
    local_report_path = None
    analysis_succeeded_and_saved = False # Flag to track if analysis part worked AND saved
    try:
        # Log task initiation
        log_progress(job_id, "init", "Task started", {
            "candidate_id": candidate_id,
            "resume_path": resume_path
            # Avoid logging potentially large JD here
        })

        # Validate job_id exists and fetch description from DB
        log_progress(job_id, "fetch_jd", f"Fetching description for job_id {job_id}")
        job_response = supabase.table("hr_jobs").select("description").eq("id", job_id).execute()
        if not job_response.data or not job_response.data[0].get("description"):
            log_progress(job_id, "error", f"Job description not found in hr_jobs for job_id {job_id}")
            raise Exception(f"Job description not found for job {job_id}")
        job_description_from_db = job_response.data[0]["description"] # Use DB description
        log_progress(job_id, "fetch_jd", "Fetched job description from DB", {"jd_length": len(job_description_from_db)})


        # Download resume
        log_progress(job_id, "download_resume", "Downloading resume")
        local_resume_path = download_resume(resume_path)
        log_progress(job_id, "download_resume", "Resume downloaded successfully", {
            "local_path": local_resume_path,
            "resume_size": os.path.getsize(local_resume_path) if local_resume_path and os.path.exists(local_resume_path) else -1
        })

        # Extract text from resume
        log_progress(job_id, "extract_text", "Extracting text from resume")
        resume_text = extract_text_from_file(local_resume_path, job_id)
        log_progress(job_id, "extract_text", "Text extracted successfully", {
            "text_length": len(resume_text)
        })

        # Generate report (Pass job_id for internal logging)
        log_progress(job_id, "generate_report", "Generating report with Gemini")
        report = generate_report(resume_text, job_description_from_db, job_id) # Pass job_id
        log_progress(job_id, "generate_report", "Report generated successfully", {
            "overall_score": report.get("overall_match_score", "N/A"),
            "candidate_name": report.get("candidate_name", "N/A")
        })

        # Ensure candidate exists in hr_candidates
        log_progress(job_id, "check_candidate", f"Checking/creating candidate {candidate_id}")
        candidate_check_resp = supabase.table("hr_candidates").select("id").eq("id", candidate_id).execute() # Select only id
        if not candidate_check_resp.data:
             log_progress(job_id, "create_candidate", f"Candidate {candidate_id} not found, creating...")
             try:
                insert_data = {
                    "id": candidate_id,
                    "name": report.get("candidate_name", "Unknown"),
                    "email": report.get("email") or f"unknown_{candidate_id}@example.com",
                    "phone_number": report.get("phone_number"), # Use .get() to allow None
                    "linkedin_url": report.get("linkedin"),
                    "github_url": report.get("github")
                }
                log_progress(job_id, "create_candidate_data", "Candidate data for insert", {"data": insert_data})
                supabase.table("hr_candidates").insert(insert_data).execute()
                log_progress(job_id, "create_candidate_success", f"Created new candidate record for {candidate_id}")
             except Exception as candidate_insert_exc:
                 log_progress(job_id, "create_candidate_error", f"Failed to insert candidate {candidate_id}: {str(candidate_insert_exc)}")
                 raise # Fail task if candidate creation fails

        # Save report as PDF
        log_progress(job_id, "save_report", "Saving report as PDF")
        report_filename = f"report_{job_id}_{candidate_id}.pdf"
        local_report_path = f"/tmp/{report_filename}"
        save_report_as_pdf(report, local_report_path, job_id) # Pass job_id
        log_progress(job_id, "save_report", "Report saved successfully", {
            "local_report_path": local_report_path
        })

        # Upload report to Supabase Storage (Pass job_id for internal logging)
        log_progress(job_id, "upload_report", "Uploading report to Supabase Storage")
        report_destination_path = f"{SUPABASE_REPORT_PATH_PREFIX}/{job_id}/{report_filename}"
        upload_report(local_report_path, report_destination_path, job_id) # Pass job_id
        log_progress(job_id, "upload_report", "Report uploaded successfully", {
            "report_destination_path": report_destination_path
        })

        # Step 1: Prepare candidate_resume_analysis payload
        log_progress(job_id, "prepare_payload", "Preparing payload for candidate_resume_analysis")
        supabase_project_id = os.getenv("SUPABASE_PROJECT_ID")
        if not supabase_project_id:
             try: supabase_project_id = SUPABASE_URL.split('.')[0].split('//')[1]
             except IndexError: supabase_project_id = "[YOUR_PROJECT_ID]" # Fallback
        report_public_url = f"https://{supabase_project_id}.supabase.co/storage/v1/object/public/{SUPABASE_STORAGE_BUCKET}/{report_destination_path}"
        resume_payload = {
            "job_id": job_id, "candidate_id": candidate_id, "resume_text": resume_text or None,
            "overall_score": round(report.get("overall_match_score", 0)),
            "matched_skills": report.get("matched_skills", []), "summary": report.get("summary"),
            "missing_or_weak_areas": report.get("missing_or_weak_areas", []),
            "top_skills": report.get("top_skills", []), "development_gaps": report.get("development_gaps", []),
            "additional_certifications": report.get("additional_certifications", []),
            "section_wise_scoring": report.get("section_wise_scoring", {}),
            "candidate_name": report.get("candidate_name", "Unknown"), "email": report.get("email", ""),
            "phone_number": report.get("phone_number", ""), "github": report.get("github", ""),
            "linkedin": report.get("linkedin", ""), "report_url": report_public_url,
            "has_validated_resume": True, # Set to True here
            "updated_at": datetime.datetime.utcnow().isoformat()
        }
        log_progress(job_id, "prepare_payload_complete", "Payload prepared")

        # Step 2: Upsert candidate_resume_analysis
        log_progress(job_id, "save_candidate_resume_analysis", "Upserting candidate_resume_analysis")
        try:
            analysis_response = supabase.table("candidate_resume_analysis").upsert(
                resume_payload, on_conflict="job_id,candidate_id"
            ).execute()
            log_progress(job_id, "save_candidate_resume_analysis_response", "Upsert response", {"data": str(analysis_response.data), "count": str(analysis_response.count)})
            analysis_succeeded_and_saved = True # Mark analysis save as successful
        except Exception as analysis_upsert_exc:
             log_progress(job_id, "save_candidate_resume_analysis_error", f"Failed to upsert candidate_resume_analysis: {str(analysis_upsert_exc)}")
             raise # Re-raise to fail the whole task

        # --- Step 3: Update hr_job_candidates (Code-based approach) ---
        if analysis_succeeded_and_saved: # Only run if previous step was definitely successful
            log_progress(job_id, "update_hr_job_candidates_start", "Attempting to update has_validated_resume in hr_job_candidates")
            try:
                # Check if the row exists before trying to update
                check_response = supabase.table("hr_job_candidates").select(
                    "job_id", count='exact' # Select minimal column, get count
                ).eq("job_id", job_id).eq("candidate_id", candidate_id).execute()

                if check_response.count == 0:
                    # Log that the row was missing - this might be expected or an error depending on your flow
                    log_progress(job_id, "update_hr_job_candidates_skip", f"Row NOT FOUND in hr_job_candidates for job {job_id}, candidate {candidate_id}. Cannot update flag.")
                else:
                    # Row exists, proceed with the update
                    log_progress(job_id, "update_hr_job_candidates_found", f"Row FOUND in hr_job_candidates. Updating has_validated_resume=True.")
                    update_response = supabase.table("hr_job_candidates").update(
                        {"has_validated_resume": True} # Set to True
                    ).eq("job_id", job_id).eq("candidate_id", candidate_id).execute()

                    # Log the response from the update attempt
                    log_progress(job_id, "update_hr_job_candidates_response", "Supabase response for hr_job_candidates update", {
                        "response_data": str(update_response.data) if hasattr(update_response, 'data') else "N/A",
                        "response_count": str(update_response.count) if hasattr(update_response, 'count') else "N/A"
                    })
                    # Assume success if execute() did not raise an error

            except Exception as update_exc:
                # Log specific errors during the check or update, but don't necessarily fail the whole task
                tb_str = traceback.format_exc()
                log_progress(job_id, "update_hr_job_candidates_error", f"Non-fatal error during hr_job_candidates check/update: {str(update_exc)}", {
                    "error_type": type(update_exc).__name__, "error_message": str(update_exc), "traceback": tb_str
                })
                logger.error(f"Job {job_id} - Non-fatal error updating hr_job_candidates: {update_exc}") # Log as error but allow task to potentially continue

        # --- Step 4: Save and check company associations ---
        log_progress(job_id, "process_companies", "Processing company associations")
        company_entries = []
        raw_companies = report.get("companies", [])
        if isinstance(raw_companies, list):
            for company in raw_companies:
                if isinstance(company, dict) and company.get("name"):
                    company_name = company["name"]
                    normalized_name = normalize_company_name(company_name)
                    if not normalized_name: continue # Skip empty normalized names
                    try:
                        company_id = None
                        company_check_response = supabase.table("companies").select("id", count='exact').eq("normalized_name", normalized_name).execute()
                        if company_check_response.count > 0:
                            company_id = company_check_response.data[0]["id"]
                        else:
                            log_progress(job_id, "process_companies_insert", f"Inserting new company: {company_name} (Normalized: {normalized_name})")
                            new_company_response = supabase.table("companies").insert({"name": company_name, "normalized_name": normalized_name}).execute()
                            if new_company_response.data:
                                company_id = new_company_response.data[0]["id"]
                            else:
                                log_progress(job_id, "process_companies_insert_error", f"Failed to insert company: {company_name}")
                                continue # Skip if insert failed
                        if company_id:
                            company_entries.append({
                                "candidate_id": candidate_id, "job_id": job_id, "company_id": company_id,
                                "designation": company.get("designation", "-"), "years": company.get("years", "-")
                            })
                    except Exception as company_lookup_exc:
                         log_progress(job_id, "process_companies_lookup_error", f"Error processing company '{company_name}': {str(company_lookup_exc)}")
            # Upsert collected company associations
            if company_entries:
                log_progress(job_id, "save_candidate_companies_start", "Upserting candidate_companies", {"count": len(company_entries)})
                try:
                    supabase.table("candidate_companies").upsert(
                        company_entries, on_conflict="candidate_id,job_id,company_id" # Adjust constraint if needed
                    ).execute()
                    log_progress(job_id, "save_candidate_companies_success", "Upserted candidate_companies successfully")
                except Exception as company_upsert_exc:
                    log_progress(job_id, "save_candidate_companies_error", f"Exception during candidate_companies upsert: {str(company_upsert_exc)}")
                    raise # Fail the whole task if this critical step fails
            else:
                 log_progress(job_id, "save_candidate_companies_skip", "No valid company entries found to upsert")
        else:
             log_progress(job_id, "process_companies_warning", "Report 'companies' field is not a list")

        # If execution reaches here, the main try block completed successfully
        log_progress(job_id, "success", "Task completed successfully")
        return {"status": "finished", "candidate_id": candidate_id}

    except Exception as e:
        # This block runs if ANY error occurred in the main try block
        tb_str = traceback.format_exc()
        logger.error(f"Job {job_id} - process_analysis_error: Task failed for candidate {candidate_id}: {str(e)}")
        logger.error(f"Job {job_id} - Traceback:\n{tb_str}")
        log_progress(job_id, "error", f"Task failed: {str(e)}", {"error_message": str(e), "traceback": tb_str})

        # --- Update status to FALSE on failure ---
        try:
            log_progress(job_id, "failure_update_status", "Attempting to set has_validated_resume=False in DB due to failure")
            # Update analysis table first (might not exist if initial steps failed, but update is safe)
            supabase.table("candidate_resume_analysis").update(
                {"has_validated_resume": False, "summary": f"Processing failed: {str(e)[:500]}"}
            ).eq("job_id", job_id).eq("candidate_id", candidate_id).execute()

            # Conditionally update job candidates table ONLY IF the row exists
            check_response = supabase.table("hr_job_candidates").select("job_id", count='exact').eq("job_id", job_id).eq("candidate_id", candidate_id).execute()
            if check_response.count > 0:
                supabase.table("hr_job_candidates").update(
                    {"has_validated_resume": False}
                ).eq("job_id", job_id).eq("candidate_id", candidate_id).execute()
                log_progress(job_id, "failure_update_hr_job_candidates_success", "Set has_validated_resume=False in hr_job_candidates")
            else:
                 log_progress(job_id, "failure_update_hr_job_candidates_skip", "Skipped setting flag in hr_job_candidates as row was not found")

            log_progress(job_id, "failure_update_status_complete", "Failure status update attempt finished")
        except Exception as db_update_e:
             # Log errors during the failure update process, but don't overwrite original error
             logger.error(f"Job {job_id} - DB_UPDATE_ERROR_ON_FAIL: Failed to update status to failed in DB: {db_update_e}")
             log_progress(job_id, "failure_update_status_error", f"Failed during failure status update: {str(db_update_e)}")
        # --- END FAILURE UPDATE ---

        # Return failure status for RQ result
        return {"status": "failed", "candidate_id": candidate_id, "error": str(e)}

    finally:
        # Cleanup logic
        try:
            log_progress(job_id, "cleanup", "Cleaning up temporary files")
            if local_resume_path and os.path.exists(local_resume_path): os.remove(local_resume_path)
            if local_report_path and os.path.exists(local_report_path): os.remove(local_report_path)
            log_progress(job_id, "cleanup", "Temporary files removed successfully")
        except Exception as cleanup_e:
            log_progress(job_id, "cleanup_error", f"Failed to clean up temporary files: {str(cleanup_e)}")

# END OF UPDATED process_analysis FUNCTION